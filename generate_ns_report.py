from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parent
TEMPLATE_PATH = ROOT / "NSProject_Report.docx"
OUTPUT_PATH = ROOT / "NSProject_Report_Filled.docx"


TITLE = "Multi-Attack Honeypot Network Security System"

PROBLEM_STATEMENT = (
    "This project addresses the need for a safe and educational network security "
    "environment in which multiple attack behaviors can be simulated, captured, "
    "classified, and visualized without exposing a real production system. The "
    "implemented platform combines honeypot services, packet-level detectors, a "
    "unified logging pipeline, machine learning based attack labeling, and an "
    "interactive dashboard so that students can study SSH brute-force behavior, "
    "port scanning, DNS misuse, ARP spoofing, and ICMP abuse in one integrated lab."
)

INTRO_PARAGRAPHS = [
    (
        "Network attacks rarely happen in isolation. In a realistic environment, an "
        "attacker may first perform reconnaissance, then probe exposed ports, attempt "
        "credential abuse, and finally escalate to protocol misuse or denial-of-service "
        "activity. This project was designed to demonstrate that broader attack chain "
        "inside a controlled mini-project environment."
    ),
    (
        "The system centers on a multi-attack honeypot manager that launches five core "
        "security services: an SSH honeypot, a port scan detector, a fake DNS resolver, "
        "an ARP spoof detector, and an ICMP detector. Each component watches a different "
        "layer or protocol family, which makes the project relevant to the layered "
        "network security concepts taught in the course."
    ),
    (
        "Captured activity is written into a unified CSV dataset so that attack events "
        "from different protocols can be analyzed on one timeline. This design supports "
        "cross-attack correlation, simplifies feature engineering, and helps students "
        "understand how logs can be transformed into a machine learning dataset."
    ),
    (
        "The machine learning pipeline uses behavioral features such as session duration, "
        "packet count, login attempts, query rate, scan speed, ARP conflict indicators, "
        "and ICMP packet characteristics. A Random Forest classifier is trained on these "
        "features to predict specific attack labels and severity classes for dashboard "
        "monitoring."
    ),
    (
        "For usability, the project also includes a modern frontend control center and a "
        "Streamlit analytics dashboard. These interfaces allow a user to start simulations, "
        "inspect captured sessions, retrain the model, and observe the overall attack "
        "distribution in a way that is easy to present during a mini-project demonstration."
    ),
]

NETWORK_DIAGRAM_LINES = [
    "Logical Network / System Diagram",
    "",
    "[Attacker / Simulation Node]",
    "SSH, DNS, ARP, ICMP, and TCP Scan traffic",
    "               |",
    "               v",
    "[HoneyPot Manager - Python Orchestrator]",
    "  |- SSH Honeypot           : TCP 2222",
    "  |- DNS Honeypot           : UDP 53",
    "  |- Port Scan Detector     : TCP connection monitoring",
    "  |- ARP Spoof Detector     : raw packet analysis",
    "  |- ICMP Detector          : ICMP packet analysis",
    "               |",
    "               v",
    "[Unified Attack Logger]",
    "data/attack_logs.csv",
    "               |",
    "               v",
    "[ML Pipeline - Feature Extraction + Random Forest]",
    "models/attack_model.pkl + label_encoder.pkl",
    "               |",
    "               v",
    "[Visualization Layer]",
    "Streamlit dashboard + React/Vite control center + local API bridge",
]

CONFIG_LINES = [
    "The implementation is software-defined, so the following logical setup was used instead of physical routers and switches alone.",
    "PC0 / Security Server - Runs Python 3.10+, honeypot_manager.py, unified logging, ML pipeline, and the integrated frontend/backend bridge.",
    "Laptop0 / Attacker Node - Used to execute attack_simulator.py, attack_simulator_portscan.py, attack_simulator_dns.py, attack_simulator_arp.py, and attack_simulator_icmp.py.",
    "Router0 / Local Lab Network - Provides isolated connectivity for the project; the system is intended only for a controlled classroom or lab environment.",
    "SSH Honeypot - Enabled on port 2222 with max_login_attempts = 3 and shell_timeout = 60 seconds.",
    "DNS Honeypot - Enabled on UDP port 53 with tunneling_subdomain_length = 50 characters and high_query_rate = 10 queries/second.",
    "Port Scan Detector - min_ports_for_scan = 5, aggressive_threshold = 10 ports/second, stealth_threshold = 1 port/second, session_timeout = 60 seconds.",
    "ARP Detector - high_arp_rate = 5 packets/second, requires administrator privileges for raw socket capture.",
    "ICMP Detector - flood_threshold = 50 packets/second, oversized_threshold = 1500 bytes, session_timeout = 60 seconds.",
    "Frontend and Dashboard - React/Vite interface runs through npm run dev:all, usually on http://localhost:8080, with the local API bridge on port 8787.",
]

RESULT_LINES = [
    "The current project dataset contains 248 captured attack sessions in the unified CSV file.",
    "Observed attack-type distribution from the available logs: SSH = 166 sessions, DNS = 73 sessions, ICMP = 4 sessions, Port Scan = 3 sessions, and ARP = 2 sessions.",
    "Generated label distribution from the present data shows ssh_normal = 148, dns_amplification = 66, ssh_reconnaissance = 18, dns_normal = 6, icmp_flood = 3, arp_poisoning = 2, portscan_sweep = 2, and single-sample cases for portscan_stealth, dns_tunneling, and icmp_ping_of_death.",
    "The project architecture supports 13 total attack labels, but the currently saved trained label encoder in the workspace includes 10 learned classes because some categories are not yet represented well in the available dataset.",
    "A single-process validation run on the current dataset using the project feature extraction pipeline and a Random Forest classifier produced 100% accuracy on an 80/20 split. This demonstrates that the pipeline is functioning correctly, but the result should be interpreted carefully because the dataset is small and imbalanced for a few attack categories.",
    "Operational proof from the source code confirms end-to-end workflow support: attack simulators generate traffic, honeypot_manager.py orchestrates the services, attack_logger.py stores unified logs, train_model_multi.py prepares the classifier, dashboard/dashboard_multi.py visualizes events, and Frontend/server/api.mjs exposes an integrated UI bridge.",
    "For the final submitted report, replace this section with 3 clear screenshots showing (1) honeypot manager running, (2) dashboard or frontend attack visualization, and (3) attack log or model training output during the live demonstration.",
]

CONCLUSION = (
    "In this mini project, a multi-attack honeypot network security platform was designed "
    "and implemented to simulate realistic attack activity across SSH, DNS, ARP, ICMP, "
    "and TCP port-scanning workflows. The system was built by combining protocol-aware "
    "detectors, a unified CSV logging layer, a Random Forest based machine learning "
    "classifier, and both dashboard and frontend monitoring interfaces. The project is "
    "useful because it gives students and defenders a safe way to study attacker behavior, "
    "logging strategy, feature engineering, and real-time security analytics in one "
    "environment. In the future, the platform can be extended with more balanced datasets, "
    "additional attack classes, stronger evaluation metrics, containerized deployment, and "
    "automated alerting for incident response workflows."
)

PHOTO_PLACEHOLDER = (
    "Insert one geo-tagged photograph of the team presentation here before final submission."
)


def paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    if style is not None:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def delete_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    paragraph._p = paragraph._element = None


def set_paragraph_text(paragraph, text, bold=False):
    paragraph.clear()
    run = paragraph.add_run(text)
    run.bold = bold
    return run


def set_font(paragraph, font_name="Times New Roman", size=12):
    for run in paragraph.runs:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
        run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
        run.font.size = Pt(size)


def ensure_spacing(paragraph, before=0, after=6, line_spacing=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing


def find_first(doc, exact_text):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == exact_text:
            return paragraph
    raise ValueError(f"Paragraph not found: {exact_text}")


def find_containing(doc, snippet):
    for paragraph in doc.paragraphs:
        if snippet in paragraph.text:
            return paragraph
    raise ValueError(f"Paragraph containing snippet not found: {snippet}")


def add_normal_paragraphs_after(anchor, texts):
    current = anchor
    for text in texts:
        current = paragraph_after(current, text)
        ensure_spacing(current)
        set_font(current)
    return current


def find_all(doc, exact_text):
    return [paragraph for paragraph in doc.paragraphs if paragraph.text.strip() == exact_text]


def find_all_containing(doc, snippet):
    return [paragraph for paragraph in doc.paragraphs if snippet in paragraph.text]


def build_report():
    doc = Document(str(TEMPLATE_PATH))

    set_paragraph_text(find_first(doc, "Title"), TITLE, bold=True)
    set_paragraph_text(find_first(doc, "Submitted"), "Submitted")
    set_paragraph_text(find_first(doc, "By"), "By")

    student_paragraphs = find_all_containing(doc, "Names")
    student_placeholders = [
        "[Student Name 1] ([USN 1])",
        "[Student Name 2] ([USN 2])",
        "[Student Name 3] ([USN 3])",
    ]
    for paragraph, text in zip(student_paragraphs, student_placeholders):
        set_paragraph_text(paragraph, text)

    set_paragraph_text(find_first(doc, "Dr./Prof."), "[Faculty Guide Name]")
    set_paragraph_text(
        find_first(
            doc,
            "Certified that the CS3403 Network Security Mini Project work titled XXXXX is carried out by (names) XXXX (USN) who are bonafide students of the School of Computer Science and Engineering, RV University, Bengaluru, during the year 2025–26. It is certified that all corrections/ suggestions from all the continuous internal evaluations have been incorporated into the project and in this report.",
        ),
        (
            "Certified that the CS3403 Network Security Mini Project work titled "
            f"\"{TITLE}\" is carried out by [Student Names] ([USNs]) who are bonafide "
            "students of the School of Computer Science and Engineering, RV University, "
            "Bengaluru, during the year 2025-26. It is certified that all corrections "
            "and suggestions from the continuous internal evaluations have been incorporated "
            "into the project and in this report."
        ),
    )
    set_paragraph_text(find_first(doc, "Dr./ Prof. ________________"), "[Faculty Guide Name]")

    problem_paragraph = find_first(doc, "One paragraph explanation on what problem/technology you are addressing/simulating")
    set_paragraph_text(problem_paragraph, PROBLEM_STATEMENT)
    ensure_spacing(problem_paragraph)
    set_font(problem_paragraph)

    intro_placeholder = find_first(
        doc,
        "Maximum of 5 paragraphs- about your network, its relevance, and a short explanation of technologies/ terminologies/ protocols involved.",
    )
    set_paragraph_text(intro_placeholder, INTRO_PARAGRAPHS[0])
    ensure_spacing(intro_placeholder)
    set_font(intro_placeholder)
    add_normal_paragraphs_after(intro_placeholder, INTRO_PARAGRAPHS[1:])

    diagram_placeholder = find_first(doc, "Paste a clear diagram of your network")
    set_paragraph_text(diagram_placeholder, "\n".join(NETWORK_DIAGRAM_LINES))
    ensure_spacing(diagram_placeholder, after=8, line_spacing=1.0)
    set_font(diagram_placeholder, font_name="Courier New", size=10)

    config_intro = find_first(doc, "Add configuration set up of each and every device used.")
    set_paragraph_text(config_intro, CONFIG_LINES[0])
    ensure_spacing(config_intro)
    set_font(config_intro)

    config_placeholders = [
        "PC0 –",
        "Laptop 0-",
        "Router 0-",
        "………..",
        "……….",
    ]
    config_targets = []
    for item in config_placeholders:
        try:
            config_targets.append(find_first(doc, item))
        except ValueError:
            pass

    for paragraph, text in zip(config_targets, CONFIG_LINES[1:1 + len(config_targets)]):
        set_paragraph_text(paragraph, text)
        ensure_spacing(paragraph)
        set_font(paragraph)

    if len(CONFIG_LINES) - 1 > len(config_targets):
        add_normal_paragraphs_after(config_targets[-1], CONFIG_LINES[1 + len(config_targets):])

    results_intro = find_containing(doc, "Paste proof of communication between devices like simulation effect/ ping status/")
    set_paragraph_text(results_intro, RESULT_LINES[0])
    ensure_spacing(results_intro)
    set_font(results_intro)
    add_normal_paragraphs_after(results_intro, RESULT_LINES[1:])

    minimum_proof = find_first(doc, "Minimum 3 proofs- should be clear screenshots.")
    delete_paragraph(minimum_proof)

    conclusion_placeholder = find_first(
        doc,
        "1 paragraph:- what you did, how you did it, how it was useful to society/network world, and what can be extended in the future.",
    )
    set_paragraph_text(conclusion_placeholder, CONCLUSION)
    ensure_spacing(conclusion_placeholder)
    set_font(conclusion_placeholder)

    photo_placeholder = find_first(doc, "Add one geotag photo of your presentation")
    set_paragraph_text(photo_placeholder, PHOTO_PLACEHOLDER)
    ensure_spacing(photo_placeholder)
    set_font(photo_placeholder)

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in {"1  Problem statement", "3  Network Diagram", "4. Configuration setup"}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if text in {"Mini Project Report", TITLE, "CERTIFICATE"}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(str(OUTPUT_PATH))


if __name__ == "__main__":
    build_report()
